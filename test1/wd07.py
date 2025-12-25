from docx import Document
doc = Document('test.docx')
p = doc.paragraphs[4]
p.add_run('문단에 글자 추가')
doc.save('test6.docx')