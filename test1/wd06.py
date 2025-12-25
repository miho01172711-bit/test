from docx import Document
doc = Document()
doc.add_heading('코딩유치원 python-docx 강의', level=0)
p = doc.add_paragraph('안녕하세요, 코린이 여러분!')
p.add_run('코딩유치원에 오신 것을 환영합니다').bold = True
doc.add_paragraph('문장 추가1')
doc.add_paragraph('문장 추가2')
doc.add_paragraph('문장 추가3')
doc.add_paragraph('문장 추가4')
records = (
    (1, '하나', 'one'),
    (2, '둘', 'two'),
    (3, '셋', 'three')
)
table = doc.add_table(rows=1, cols=3)
table.style = doc.styles['Table Grid']
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = '한국어'
hdr_cells[2].text = '영어'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

doc.save('test.docx')