from docx import Document
doc = Document()
doc.add_paragraph('첫 번째 문단입니다.')
doc.add_paragraph('글머리 문단', style='List Bullet')
doc.add_paragraph('글머리 문단', style='List Bullet')
doc.add_paragraph('글머리 문단', style='List Bullet')
doc.add_paragraph('번호 목록', style='List Number')
doc.add_paragraph('번호 목록', style='List Number')
doc.add_paragraph('번호 목록', style='List Number')
p = doc.add_paragraph('두 번째 문단입니다.')
p.add_run('문단에 굵은 글짜 추가').bold = True
p.add_run('이탤릭체').italic = True
p.add_run('밑줄인 글자').underline = True
doc.save('test3.docx')