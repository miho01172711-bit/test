from docx import Document
doc = Document("test.docx")
tables = doc.tables

print(tables[0].rows[0].cells[0].paragraphs[0].text)
table = tables[0]
for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            print(para.text, end=' ')
    print()