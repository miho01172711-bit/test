from docx import Document
from docx.shared import Cm, Inches
doc = Document()
doc.add_picture('dog.jpg', width=Cm(5), height=Cm(5))
doc.add_picture('dog.jpg', width=Inches(4), height=Inches(3))
doc.save('test4.docx')