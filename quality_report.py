import os
from datetime import datetime

import pandas as pd
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt


# =========================
# 설정
# =========================
INPUT_DIR = "out_csv"
OUTPUT_DIR = "out"

SUMMARY_XLSX = os.path.join(OUTPUT_DIR, "quality_summary.xlsx")
REPORT_DOCX = os.path.join(OUTPUT_DIR, "quality_report.docx")

CSV_FILES = {
    "summary": "summary.csv",
    "tests": "tests.csv",
    "coverage_files": "coverage_files.csv",
    "lint": "lint.csv",
    "lint_top5": "lint_top5.csv",
}


# =========================
# 유틸
# =========================
def _require_file(path: str) -> None:
    if not os.path.exists(path):
        raise FileNotFoundError(f"필수 입력 파일이 없습니다: {path}")


def _safe_float(x, default=0.0) -> float:
    try:
        return float(x)
    except Exception:
        return default


def _auto_fit_columns(ws) -> None:
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            v = "" if cell.value is None else str(cell.value)
            if len(v) > max_len:
                max_len = len(v)
        ws.column_dimensions[col_letter].width = min(max_len + 2, 60)


def _add_table_docx(doc: Document, title: str, df: pd.DataFrame, max_rows: int = 30) -> None:
    doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("(데이터 없음)")
        return

    show_df = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(show_df.columns))
    table.style = "Table Grid"

    # header
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(show_df.columns):
        hdr_cells[i].text = str(c)

    # rows
    for _, row in show_df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(show_df.columns):
            cells[i].text = "" if pd.isna(row[c]) else str(row[c])

    if len(df) > max_rows:
        doc.add_paragraph(f"(표는 상위 {max_rows}행만 표시됨 / 전체 {len(df)}행)")


# =========================
# 메인 로직
# =========================
def main() -> None:
    # 1) 입력 파일 존재 확인
    paths = {k: os.path.join(INPUT_DIR, v) for k, v in CSV_FILES.items()}
    for p in paths.values():
        _require_file(p)

    # 2) CSV 읽기
    summary_df = pd.read_csv(paths["summary"])
    tests_df = pd.read_csv(paths["tests"])
    coverage_df = pd.read_csv(paths["coverage_files"])
    lint_df = pd.read_csv(paths["lint"])
    lint_top5_df = pd.read_csv(paths["lint_top5"])

    # 3) summary.csv를 dict로 변환(metric,value)
    #    예: tests_total, tests_passed, coverage_percent, lint_issues ...
    summary_map = {}
    if not summary_df.empty and {"metric", "value"}.issubset(summary_df.columns):
        for _, r in summary_df.iterrows():
            summary_map[str(r["metric"])] = r["value"]

    tests_total = int(_safe_float(summary_map.get("tests_total", len(tests_df))))
    tests_passed = int(_safe_float(summary_map.get("tests_passed", 0)))
    tests_failed = int(_safe_float(summary_map.get("tests_failed", 0)))
    tests_skipped = int(_safe_float(summary_map.get("tests_skipped", 0)))
    tests_time_sec = _safe_float(summary_map.get("tests_time_sec", 0.0))

    coverage_percent = _safe_float(summary_map.get("coverage_percent", 0.0))
    lint_issues = int(_safe_float(summary_map.get("lint_issues", len(lint_df))))

    # 4) out 폴더 생성
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # =========================
    # A) Excel 생성: out/quality_summary.xlsx
    # - Summary(원본 metric/value)
    # - Tests(테스트 케이스)
    # - Coverage(파일별 커버리지)
    # - Lint(린트 상세)
    # - LintTop5(Top5)
    # =========================
    wb = Workbook()

    # Sheet 1: Summary
    ws = wb.active
    ws.title = "Summary"
    ws.append(["metric", "value"])
    for _, r in summary_df.iterrows():
        ws.append([r.get("metric", ""), r.get("value", "")])
    _auto_fit_columns(ws)

    # Sheet 2: Tests
    ws2 = wb.create_sheet("Tests")
    if not tests_df.empty:
        ws2.append(list(tests_df.columns))
        for _, r in tests_df.iterrows():
            ws2.append([("" if pd.isna(v) else v) for v in r.tolist()])
    else:
        ws2.append(["(no data)"])
    _auto_fit_columns(ws2)

    # Sheet 3: Coverage
    ws3 = wb.create_sheet("Coverage")
    if not coverage_df.empty:
        ws3.append(list(coverage_df.columns))
        for _, r in coverage_df.iterrows():
            ws3.append([("" if pd.isna(v) else v) for v in r.tolist()])
    else:
        ws3.append(["(no data)"])
    _auto_fit_columns(ws3)

    # Sheet 4: Lint
    ws4 = wb.create_sheet("Lint")
    if not lint_df.empty:
        ws4.append(list(lint_df.columns))
        for _, r in lint_df.iterrows():
            ws4.append([("" if pd.isna(v) else v) for v in r.tolist()])
    else:
        ws4.append(["(no data)"])
    _auto_fit_columns(ws4)

    # Sheet 5: LintTop5
    ws5 = wb.create_sheet("LintTop5")
    if not lint_top5_df.empty:
        ws5.append(list(lint_top5_df.columns))
        for _, r in lint_top5_df.iterrows():
            ws5.append([("" if pd.isna(v) else v) for v in r.tolist()])
    else:
        ws5.append(["(no data)"])
    _auto_fit_columns(ws5)

    wb.save(SUMMARY_XLSX)

    # =========================
    # B) Word 생성: out/quality_report.docx
    # - 요약(핵심 지표)
    # - 테스트 요약 + 실패 케이스 상위 표시
    # - 커버리지 Top/Bottom
    # - 린트 Top5 + 상세 일부
    # =========================
    doc = Document()

    # 기본 글꼴 크기(선택)
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10)

    doc.add_heading("Software Quality Report", level=1)
    doc.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 1) 핵심 요약
    doc.add_heading("1. Executive Summary", level=2)
    doc.add_paragraph(f"- Total tests: {tests_total}")
    doc.add_paragraph(f"- Passed: {tests_passed}, Failed: {tests_failed}, Skipped: {tests_skipped}")
    doc.add_paragraph(f"- Test time (sec): {tests_time_sec}")
    doc.add_paragraph(f"- Coverage (%): {coverage_percent}")
    doc.add_paragraph(f"- Lint issues: {lint_issues}")

    # 2) Summary 원본 표
    _add_table_docx(doc, "2. Raw Summary (summary.csv)", summary_df, max_rows=50)

    # 3) Tests 결과 (실패/스킵 케이스 일부)
    doc.add_heading("3. Test Results (tests.csv)", level=2)
    if tests_df.empty:
        doc.add_paragraph("(테스트 데이터 없음)")
    else:
        doc.add_paragraph(f"Total testcases: {len(tests_df)}")

        # 실패/스킵 케이스 우선 표시
        if "status" in tests_df.columns:
            failed = tests_df[tests_df["status"].astype(str).str.lower() == "failed"]
            skipped = tests_df[tests_df["status"].astype(str).str.lower() == "skipped"]

            if not failed.empty:
                _add_table_docx(doc, "3.1 Failed Testcases (top 20)", failed, max_rows=20)
            else:
                doc.add_paragraph("3.1 Failed Testcases: (none)")

            if not skipped.empty:
                _add_table_docx(doc, "3.2 Skipped Testcases (top 20)", skipped, max_rows=20)
            else:
                doc.add_paragraph("3.2 Skipped Testcases: (none)")
        else:
            _add_table_docx(doc, "3.1 Testcases (top 30)", tests_df, max_rows=30)

    # 4) Coverage 요약
    doc.add_heading("4. Coverage (coverage_files.csv)", level=2)
    if coverage_df.empty:
        doc.add_paragraph("(커버리지 데이터 없음)")
    else:
        if "coverage_percent" in coverage_df.columns:
            cov_sorted = coverage_df.copy()
            cov_sorted["coverage_percent"] = cov_sorted["coverage_percent"].apply(_safe_float)

            top5 = cov_sorted.sort_values("coverage_percent", ascending=False).head(5)
            bottom5 = cov_sorted.sort_values("coverage_percent", ascending=True).head(5)

            _add_table_docx(doc, "4.1 Top 5 Coverage Files", top5, max_rows=5)
            _add_table_docx(doc, "4.2 Bottom 5 Coverage Files", bottom5, max_rows=5)
        else:
            _add_table_docx(doc, "4.1 Coverage Files (top 30)", coverage_df, max_rows=30)

    # 5) Lint 요약
    doc.add_heading("5. Lint (ruff) Results", level=2)
    if lint_top5_df.empty:
        doc.add_paragraph("(lint_top5 데이터 없음)")
    else:
        _add_table_docx(doc, "5.1 Top 5 Lint Rules", lint_top5_df, max_rows=10)

    if lint_df.empty:
        doc.add_paragraph("5.2 Lint details: (none)")
    else:
        _add_table_docx(doc, "5.2 Lint details (top 30)", lint_df, max_rows=30)

    # 6) 결론
    doc.add_heading("6. Conclusion", level=2)
    doc.add_paragraph(
        "본 보고서는 테스트(pytest), 커버리지(coverage), 정적분석(ruff) 결과를 CSV로 집계한 뒤, "
        "Excel과 Word 산출물로 자동 생성한다. 이를 통해 CI 환경에서 품질 지표를 일관되게 확인하고 "
        "보고서 형태로 공유할 수 있다."
    )

    doc.save(REPORT_DOCX)

    print("✔ 완료")
    print(f"- Excel: {SUMMARY_XLSX}")
    print(f"- Word : {REPORT_DOCX}")


if __name__ == "__main__":
    main()
