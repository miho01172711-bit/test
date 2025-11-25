from docx import Document
from docx.shared import Pt

doc = Document()
para = doc.add_paragraph('이 글자의 크기를 바꿔봅시다.')
para = doc.paragraphs[0].runs
para[0].font.size = Pt(20)
doc.save('test12.docx')